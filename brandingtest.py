brands = ['Android ','Chrome','Adobe','Acrobat','Elasticsearch','VirtualBox','AnyDesk','Apache','Apple','Asterisk','Atlasian','Atlassian','Avaya','Avira','BankIt','isc-Bind','Bitedefender','Canon','CentOs','Checkpoint','Cisco','Citrix','Cpanel','Cyberoam','Debian','dell','Diebold Nixdorf','Dlink','Docker','DotNetNuke','Drupal','EMC',' ESET','F5','Fedoraproject','Forti','FreeBSD','F-ssecure','Google','GRG','GP','Huawei','IBM','Ingenico','Intel','Issabel','Java','Jenkins','Jetbrains','Joomla','Juniper','Kaspersky','Kayako','Kerio','Kubernetes','Linux','Mcafee','Microsoft','Mikrotik','Mongodb','Mozilla','Mysql','NCR','Nginx','Norton','Nvidia','Omron','Opensuse','Oracle','OTRS','Palo Alto','Paessler','pfSense','PHP','PostgreSQL','Prometheus','PRTG','Python','Qemu','QNAP','Qualcomm','Redhat','Redis','RTIR','Samsung','SAP ','Schneider','Solaris','Solarwinds','Sonicwall','Sophos','Splunk','Symantec','Teamviewer','Tp-link','Ubuntu','Verifone','Vmware','Wincor','Wireshark','WordPress','Zabbix','Zoho']
from typing import Counter
import pandas as pd
import docx, os
def diffList(list1, list2):     # returns the difference between two lists.
    if len(list1) > len(list2):
        return (list(set(list1) - set(list2)))
    else:
        return (list(set(list2) - set(list1)))

def add_hyperlink(paragraph, url, text, color, underline): # returns the hyper link.
       
    
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')

    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is neede (kosbazi)
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
ub= [] #used brands
ur = [] #used rows 
cub = [] #cleared used brands
df = pd.read_csv('/Users/mojtaba/Desktop/file.csv')
yek2len = [] #a list that contains numbers from 1 to len(ur)
for each in brands:
    for harkudum in df['توضیحات']:
        if each.lower() in harkudum.lower():
            ub.append(each.lower())
    
for each in brands:
    if each.lower() in ub:
        cub.append(each.lower())
os.system('mkdir /Users/mojtaba/Desktop/CVES')
os.system('mkdir /Users/mojtaba/Desktop/CVES/other')

for each in cub:
    doc = docx.Document()
    doc.save('/Users/mojtaba/Desktop/CVES/%s'%(each +'.docx'))
for each in cub:
    counter = -1
    for harkudum in df['توضیحات']:
        counter +=1
        if each in harkudum.lower():
            doc = docx.Document('/Users/mojtaba/Desktop/CVES/%s'%(each + '.docx'))
            doc.add_paragraph(str(counter))
            p = doc.add_paragraph()
            hyperlink = add_hyperlink(p, df['لینک'][counter], df['سریال'][counter], 'FF8822', False)
            doc.add_paragraph(df['توضیحات'][counter])
            doc.add_paragraph('nvd = ' +df['امتیاز ان وی دی'][counter] )
            doc.add_paragraph('Vendor = '+df['امتیاز وندور '][counter])
            b = doc.add_paragraph()
            hyperlink = add_hyperlink(b, df['راهکار'][counter], 'راهکار', 'FF8822', False)
            doc.add_paragraph('---------')
            doc.save('/Users/mojtaba/Desktop/CVES/%s'%(each + '.docx'))
            ur.append(int(counter))
ur = sorted(list(dict.fromkeys(ur)))
for each in range(0,int(len(df['توضیحات'])+1)):
    yek2len.append(each)
difr = diffList(yek2len,ur)
doc = docx.Document()
doc.save('/Users/mojtaba/Desktop/CVES/other/not_related.docx')
for each in difr:
    if each== len(df['توضیحات']):
        pass
    else:
        doc = docx.Document('/Users/mojtaba/Desktop/CVES/other/not_related.docx')
        doc.add_paragraph(str(each))
        p = doc.add_paragraph()
        hyperlink = add_hyperlink(p, df.loc[each][5], df.loc[each][0], 'FF8822', False)
        doc.add_paragraph(df.loc[each][4])
        doc.add_paragraph('nvd = ' +df.loc[each][2] )
        doc.add_paragraph('Vendor = '+df.loc[each][1])
        b = doc.add_paragraph()
        hyperlink = add_hyperlink(b, df.loc[each][3], 'راهکار', 'FF8822', False)
        doc.add_paragraph('---------')
        doc.save('/Users/mojtaba/Desktop/CVES/other/not_related.docx')



