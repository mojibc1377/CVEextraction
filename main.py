import requests, bs4, re, csv, time
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from requests.api import head
from typing import Counter
import pandas as pd
import docx, os
brands = ['Android ','Chrome','Adobe','Elasticsearch','VirtualBox','AnyDesk','Apache','Apple','Asterisk','Atlassian','Avaya','Avira','BankIt','isc-Bind','Bitedefender','Canon','CentOs','Checkpoint','Cisco','Citrix','Cpanel','Cyberoam','Debian','dell','Diebold Nixdorf','Dlink','Docker','DotNetNuke','Drupal','EMC',' ESET','F5','Fedoraproject','Forti','FreeBSD','F-ssecure','Google','GRG','GP','Huawei','IBM','Ingenico','Intel','Issabel','Java','Jenkins','Jetbrains','Joomla','Juniper','Kaspersky','Kayako','Kerio','Kubernetes','Linux','Mcafee','Microsoft','Mikrotik','Mongodb','Mozilla','Mysql','NCR','Nginx','Norton','Nvidia','Omron','Opensuse','Oracle','OTRS','Palo Alto','Paessler','pfSense','PHP','PostgreSQL','Prometheus','PRTG','Python','Qemu','QNAP','Qualcomm','Redhat','Redis','RTIR','Samsung','SAP ','Schneider','Solaris','Solarwinds','Sonicwall','Sophos','Splunk','Symantec','Teamviewer','Tp-link','Ubuntu','Verifone','Vmware','Wincor','Wireshark','WordPress','Zabbix','Zoho']
dic ={
    '1'  : '۱',
    '1.0':'۱',
    '1.1': '۱.۱',
    '1.2': '۱.۲',
    '1.3': '۱.۳',
    '1.4': '۱.۴',
    '1.5': '۱.۵',
    '1.6': '۱.۶',
    '1.7': '۱.۷',
    '1.8': '۱.۸',
    '1.9': '۱.۹',
    '2'  : '۲',
    '2.0':'۲',
    '2.1': '۲.۱',
    '2.2': '۲.۲',
    '2.3': '۲.۳',
    '2.4': '۲.۴',
    '2.5': '۲.۵',
    '2.6': '۲.۶',
    '2.7': '۲.۷',
    '2.8': '۲.۸',
    '2.9': '۲.۹',
    '3'  : '۳',
    '3.0':'۳',
    '3.1': '۳.۱',
    '3.2': '۳.۲',
    '3.3': '۳.۳',
    '3.4': '۳.۴',
    '3.5': '۳.۵',
    '3.6': '۳.۶',
    '3.7': '۳.۷',
    '3.8': '۳.۸',
    '3.9': '۳.۹',
    '4'  : '۴',
    '4.0':'۴',
    '4.1': '۴.۱',
    '4.2': '۴.۲',
    '4.3': '۴.۳',
    '4.4': '۴.۴',
    '4.5': '۴.۵',
    '4.6': '۴.۶',
    '4.7': '۴.۷',
    '4.8': '۴.۸',
    '4.9': '۴.۹',
    '5'  : '۵',
    '5.0':'۵',
    '5.1': '۵.۱',
    '5.2': '۵.۲',
    '5.3': '۵.۳',
    '5.4': '۵.۴',
    '5.5': '۵.۵',
    '5.6': '۵.۶',
    '5.7': '۵.۷',
    '5.8': '۵.۸',
    '5.9': '۵.۹',
    '6'  : '۶',
    '6.0':'۶',
    '6.1': '۶.۱',
    '6.2': '۶.۲',
    '6.3': '۶.۳',
    '6.4': '۶.۴',
    '6.5': '۶.۵',
    '6.6': '۶.۶',
    '6.7': '۶.۷',
    '6.8': '۶.۸',
    '6.9': '۶.۹',
    '7'  : '۷',
    '7.0':'۷',
    '7.1': '۷.۱',
    '7.2': '۷.۲',
    '7.3': '۷.۳',
    '7.4': '۷.۴',
    '7.5': '۷.۵',
    '7.6': '۷.۶',
    '7.7': '۷.۷',
    '7.8': '۷.۸',
    '7.9': '۷.۹',
    '8'  : '۸',
    '8.0':'۸',
    '8.1': '۸.۱',
    '8.2': '۸.۲',
    '8.3': '۸.۳',
    '8.4': '۸.۴',
    '8.5': '۸.۵',
    '8.6': '۸.۶',
    '8.7': '۸.۷',
    '8.8': '۸.۸',
    '8.9': '۸.۹',
    '9'  : '۹',
    '9.0':'۹',
    '9.1': '۹.۱',
    '9.2': '۹.۲',
    '9.3': '۹.۳',
    '9.4': '۹.۴',
    '9.5': '۹.۵',
    '9.6': '۹.۶',
    '9.7': '۹.۷',
    '9.8': '۹.۸',
    '9.9': '۹.۹',
    '10.0': '۱۰',
    'CRITICAL' : 'بحرانی',
    'HIGH':'بالا',
    'LOW':'پایین',
    'MEDIUM':'متوسط',
}

def diffList(list1, list2):     # returns the difference between two lists.
    if len(list1) > len(list2):
        return (list(set(list1) - set(list2)))
    else:
        return (list(set(list2) - set(list1)))
def add_hyperlink(paragraph, url, text, color, underline): # returns the hyper link.
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Add color if it is needed (kosbazi)
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink 
print('yyyy/mm' )
date = input('Insert date plz:'+ '\n')
a = time.ctime()
print(a[11:19])
linkekolli = [] #list e tamame linkaye un safe
http_proxy  = "http://10.10.1.10:3128"
https_proxy = "https://10.10.1.11:1080"
ftp_proxy   = "ftp://10.10.1.10:3128"
proxyDict = { 
              "http"  : http_proxy, 
              "https" : https_proxy, 
              "ftp"   : ftp_proxy
            }
neshani = 'https://nvd.nist.gov/vuln/full-listing/' + str(date)
r1 = requests.get(neshani)
soupekoll = bs4.BeautifulSoup(r1.text,'html.parser')  
y = soupekoll.find_all('span', attrs={'class': "col-md-2"})
file = open('/Users/mojtaba/Desktop/file.csv', 'w')
header = ['سریال', 'امتیاز وندور ', 'امتیاز ان وی دی', 'راهکار', 'توضیحات', 'لینک']
writer = csv.writer(file)
writer.writerow(header)

for each in y:
    printer = re.search(r'href=\"(.*)?\"', str(each)).group(1)
    linkekolli.append('https://nvd.nist.gov/' + str(printer))
for each in linkekolli:
    address = each
    r = requests.get(address)
    sol = []
    soup = bs4.BeautifulSoup(r.text, 'html.parser')
    desc = soup.find_all('p', attrs={'data-testid': 'vuln-description'})
    desc1 = re.search(r'>(.*)<', str(desc))
    desc = desc1.group(1)
    for i in range(0, 2):
        x = soup.find_all('td', attrs={'data-testid': "vuln-hyperlinks-link-%i" % i})
        if x != []:
            x = re.search(r'href=(.*)target', str(x)).group(1)
            sol.append(x[0:-1])
        else:
            x = 'None'
            pass
    
    serial = soup.find('span', attrs={'data-testid': "page-header-vuln-id"})
    serial1 = re.search(r'>(.*)<', str(serial))
    serial = serial1.group(0)
    nvd_score = soup.find_all('a', attrs={'data-testid': "vuln-cvss3-panel-score"})
    nvd1 = re.search(r'>(.*)\<', str(nvd_score))
    if nvd1:
        nvd_score = nvd1.group(0)
    vendor_score = soup.find_all('a', attrs={'data-testid': "vuln-cvss3-cna-panel-score"})
    vector1 = re.search(r'>(.*)\<', str(vendor_score))
    if vector1:
        vendor_score = vector1.group(0)
    if nvd1:
        nvd_score = nvd_score[1:-1]
        nf = re.search(r'(.*)\s',nvd_score)
        ns = re.search(r'\s(.+)',nvd_score)
        ns = ns.group(1)
        nf = nf.group(1)
        nvd_score = dic[nf] +' , '+ dic[ns]
    else:
        nvd_score = ' نامشخص '
    if vector1:
        vf = re.search(r'(.*)\s',vendor_score[1:-1])
        vs = re.search(r'\s(.+)',vendor_score[1:-1])
        vs = vs.group(1)
        vf = vf.group(1)
        vendor_score = dic[vf] +' , '+ dic[vs]

    else:
        vendor_score = ' نامشخص '

    for each in sol:
        solution_link = str(each[1:-1])
        
    data = [ serial[1:-1], vendor_score, nvd_score, solution_link, desc[1:-1], address]
    writer.writerow(data)
file.close()
b = time.ctime()
print(b[11:19])

ub= [] #used brands
ur = [] #used rows 
cub = [] #cleared used brands
df = pd.read_csv('/Users/mojtaba/Desktop/file.csv')
yek2len = [] 
for each in brands:
    for harkudum in df['توضیحات']:
        if each.lower() in harkudum.lower():
            ub.append(each.lower())
    
for each in brands:
    if each.lower() in ub:
        cub.append(each.lower())
      
os.system('mkdir /Users/mojtaba/Desktop/CVES')
os.system('mkdir /Users/mojtaba/Desktop/CVES/other')

for each in cub:
    doc = docx.Document()
    doc.save('/Users/mojtaba/Desktop/CVES/%s'%(each +'.docx'))#TODO
for each in cub:
    counter = -1
    for harkudum in df['توضیحات']:
        counter +=1
        if each in harkudum.lower():
            doc = docx.Document('/Users/mojtaba/Desktop/CVES/%s'%(each + '.docx'))#TODO
            doc.add_paragraph(str(counter))
            p = doc.add_paragraph()
            hyperlink = add_hyperlink(p, df['لینک'][counter], df['سریال'][counter], 'FF8822', False)
            doc.add_paragraph(df['توضیحات'][counter])
            doc.add_paragraph('nvd = ' +df['امتیاز ان وی دی'][counter] )
            doc.add_paragraph('Vendor = '+df['امتیاز وندور '][counter])
            b = doc.add_paragraph()
            hyperlink = add_hyperlink(b, df['راهکار'][counter], 'راهکار', 'FF8822', False)
            doc.add_paragraph('---------')
            doc.save('/Users/mojtaba/Desktop/CVES/%s'%(each + '.docx'))#TODO
            ur.append(int(counter))

ur = sorted(list(dict.fromkeys(ur)))
for each in range(0,int(len(df['توضیحات'])+1)):
    yek2len.append(each)
difr = diffList(yek2len,ur)
doc = docx.Document()
doc.save('/Users/mojtaba/Desktop/CVES/other/not_related.docx')#TODO
for each in difr:
    if each== len(df['توضیحات']):
        pass
    else:
        doc = docx.Document('/Users/mojtaba/Desktop/CVES/other/not_related.docx') #TODO
        doc.add_paragraph(str(each))
        p = doc.add_paragraph()
        hyperlink = add_hyperlink(p, df.loc[each][5], df.loc[each][0], 'FF8822', False)
        doc.add_paragraph(df.loc[each][4])
        doc.add_paragraph('nvd = ' +df.loc[each][2] )
        doc.add_paragraph('Vendor = '+df.loc[each][1])
        b = doc.add_paragraph()
        hyperlink = add_hyperlink(b, df.loc[each][3], 'راهکار', 'FF8822', False)
        doc.add_paragraph('---------')
        doc.save('/Users/mojtaba/Desktop/CVES/other/not_related.docx')#TODO

